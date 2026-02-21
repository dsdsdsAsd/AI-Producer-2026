"""
Загрузка и обработка документов различных форматов.
Поддерживает: PDF, DOCX, TXT.
"""

from pathlib import Path
from typing import List, Tuple
import pypdf
import docx

from utils.logger import logger


class DocumentLoader:
    """Загрузчик документов различных форматов"""
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt"}
    
    @staticmethod
    def load_pdf(file_path: Path) -> str:
        """
        Загрузить PDF файл.
        
        Args:
            file_path: Путь к PDF файлу
            
        Returns:
            str: Извлеченный текст
        """
        logger.info(f"Loading PDF: {file_path}")
        
        text_parts = []
        
        try:
            with open(file_path, "rb") as file:
                pdf_reader = pypdf.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                
                logger.info(f"PDF has {total_pages} pages")
                
                for page_num, page in enumerate(pdf_reader.pages, start=1):
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        # Добавляем маркер страницы для метаданных
                        text_parts.append(f"\n[Page {page_num}]\n{page_text}")
                
                full_text = "\n".join(text_parts)
                logger.info(f"Extracted {len(full_text)} characters from PDF")
                
                return full_text
                
        except Exception as e:
            logger.error(f"Error loading PDF {file_path}: {e}")
            raise
    
    @staticmethod
    def load_docx(file_path: Path) -> str:
        """
        Загрузить DOCX файл.
        
        Args:
            file_path: Путь к DOCX файлу
            
        Returns:
            str: Извлеченный текст
        """
        logger.info(f"Loading DOCX: {file_path}")
        
        try:
            doc = docx.Document(file_path)
            
            # Извлекаем текст из параграфов
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Извлекаем текст из таблиц
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        table_texts.append(row_text)
            
            # Объединяем все
            full_text = "\n\n".join(paragraphs)
            
            if table_texts:
                full_text += "\n\n[Tables]\n" + "\n".join(table_texts)
            
            logger.info(f"Extracted {len(full_text)} characters from DOCX")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error loading DOCX {file_path}: {e}")
            raise
    
    @staticmethod
    def load_txt(file_path: Path) -> str:
        """
        Загрузить TXT файл.
        
        Args:
            file_path: Путь к TXT файлу
            
        Returns:
            str: Содержимое файла
        """
        logger.info(f"Loading TXT: {file_path}")
        
        try:
            # Пробуем разные кодировки
            encodings = ["utf-8", "cp1251", "latin-1"]
            
            for encoding in encodings:
                try:
                    with open(file_path, "r", encoding=encoding) as file:
                        text = file.read()
                    
                    logger.info(f"Loaded TXT with {encoding} encoding: {len(text)} characters")
                    return text
                    
                except UnicodeDecodeError:
                    continue
            
            # Если ни одна кодировка не подошла
            raise ValueError(f"Could not decode file {file_path} with any encoding")
            
        except Exception as e:
            logger.error(f"Error loading TXT {file_path}: {e}")
            raise
    
    @classmethod
    def load_document(cls, file_path: Path | str) -> Tuple[str, str]:
        """
        Загрузить документ (автоопределение формата).
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            Tuple[str, str]: (текст, тип файла)
            
        Raises:
            ValueError: Если формат не поддерживается
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.suffix.lower()
        
        if extension not in cls.SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file format: {extension}. "
                f"Supported: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
            )
        
        # Загружаем в зависимости от типа
        if extension == ".pdf":
            text = cls.load_pdf(file_path)
        elif extension == ".docx":
            text = cls.load_docx(file_path)
        elif extension == ".txt":
            text = cls.load_txt(file_path)
        else:
            raise ValueError(f"Unsupported extension: {extension}")
        
        return text, extension
    
    @classmethod
    def is_supported(cls, file_path: Path | str) -> bool:
        """
        Проверить, поддерживается ли формат файла.
        
        Args:
            file_path: Путь к файлу
            
        Returns:
            bool: True если поддерживается
        """
        extension = Path(file_path).suffix.lower()
        return extension in cls.SUPPORTED_EXTENSIONS


# Удобные функции для прямого использования
def load_document(file_path: Path | str) -> Tuple[str, str]:
    """
    Загрузить документ.
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        Tuple[str, str]: (текст, тип файла)
    """
    return DocumentLoader.load_document(file_path)


def is_supported_format(file_path: Path | str) -> bool:
    """
    Проверить поддержку формата.
    
    Args:
        file_path: Путь к файлу
        
    Returns:
        bool: True если поддерживается
    """
    return DocumentLoader.is_supported(file_path)
